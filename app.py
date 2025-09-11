import os
import re
import json
import uuid
import tempfile
from pathlib import Path
from flask import Flask, render_template, request, redirect, url_for, send_file, flash
from docx import Document
import pdfplumber
import requests

import nltk
nltk.download('wordnet', quiet=True)
nltk.download('omw-1.4', quiet=True)
from nltk.stem import WordNetLemmatizer
import numpy as np

app = Flask(__name__)
app.secret_key = os.getenv("FLASK_SECRET", "devkey")
GENERATED_DIR = Path("generated")
GENERATED_DIR.mkdir(exist_ok=True)

lemmatizer = WordNetLemmatizer()

STOPWORDS = {
    "and","or","the","a","an","to","for","in","on","with","by","of","is","are",
    "that","this","as","be","has","have","at","from","will","it","its","i","my"
}

# -------------------------
# Text extraction
# -------------------------
def extract_text_from_pdf(file_path):
    text = ""
    try:
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                text += (page.extract_text() or "") + "\n"
    except Exception as e:
        print("pdfplumber error:", e)
    return text.strip()

def extract_text_from_docx(file_path):
    text = ""
    try:
        doc = Document(file_path)
        for para in doc.paragraphs:
            text += para.text + "\n"
    except Exception as e:
        print("docx extraction error:", e)
    return text.strip()

def tokenize(text):
    tokens = re.findall(r'\b\w+\b', (text or "").lower())
    tokens = [lemmatizer.lemmatize(t) for t in tokens if t not in STOPWORDS and len(t) > 1]
    return tokens

# -------------------------
# Compute keyword match
# -------------------------
def compute_match_score(resume_text, job_desc):
    job_tokens = set(tokenize(job_desc))
    resume_tokens = set(tokenize(resume_text))
    common = sorted(list(job_tokens.intersection(resume_tokens)))
    missing = sorted(list(job_tokens - resume_tokens))

    if not job_tokens:
        return 0.0, common, missing
    score = (len(common) / len(job_tokens)) * 100.0
    return round(score, 2), common, missing

# -------------------------
# Ollama GPT-OSS call
# -------------------------
def call_ollama_optimize(resume_text, job_desc):
    url = "http://localhost:11434/api/generate"

    system_prompt = (
        "You are an expert career coach and professional resume writer. "
        "You will receive a candidate's raw resume text and a target job description. "
        "Your task is to COMPLETELY REWRITE the resume from scratch, tailored specifically to the job description. "
        "Do NOT reuse or repeat sentences verbatim from the input resume. "
        "Use professional resume formatting, strong action verbs, and measurable achievements. "
        "Incorporate relevant skills and keywords from the job description naturally. "
        "Highlight achievements using numbers, percentages, or outcomes wherever possible. "
        "Return JSON ONLY with keys: optimized_resume, suggested_improvements, roadmap, matched_keywords."
    )

    user_prompt = f"Resume:\n{resume_text}\n\nTarget job description:\n{job_desc}\n\nReturn JSON ONLY."

    payload = {
        "model": "gpt-oss:20b",
        "prompt": f"{system_prompt}\n\n{user_prompt}",
        "stream": False
    }

    try:
        resp = requests.post(url, json=payload)
        if resp.status_code != 200:
            print("Ollama error:", resp.text)
            return fallback_resume_result(resume_text, job_desc)

        content = resp.json().get("response", "").strip()
        print("DEBUG RAW AI RESPONSE:", content[:500])

        # Remove markdown fences like ```json
        content = re.sub(r"^```[a-zA-Z]*|```$", "", content, flags=re.MULTILINE).strip()

        # Extract JSON substring safely
        start, end = content.find("{"), content.rfind("}")
        if start != -1 and end != -1:
            content = content[start:end+1]

        try:
            data = json.loads(content)
        except Exception as e:
            print("JSON parsing failed:", e)
            return fallback_resume_result(resume_text, job_desc)

        # Ensure keys exist
        data.setdefault("optimized_resume", resume_text)
        data.setdefault("suggested_improvements", [])
        data.setdefault("roadmap", [])
        data.setdefault("matched_keywords", sorted(
            list(set(tokenize(resume_text)).intersection(set(tokenize(job_desc))))
        ))

        return data

    except Exception as e:
        print("Ollama call failed:", e)
        return fallback_resume_result(resume_text, job_desc)


def fallback_resume_result(resume_text, job_desc):
    return {
        "optimized_resume": resume_text,
        "suggested_improvements": ["(Fallback) AI request failed; review manually."],
        "roadmap": [{"step": "Manual review", "time_estimate": "1 week"}],
        "matched_keywords": sorted(list(set(tokenize(resume_text)).intersection(set(tokenize(job_desc)))))
    }

# -------------------------
# Save DOCX
# -------------------------
def save_docx_from_text(text, filename_path):
    doc = Document()
    for line in text.splitlines():
        doc.add_paragraph(line)
    doc.save(filename_path)

# -------------------------
# Routes
# -------------------------
@app.route("/", methods=["GET"])
def index():
    return render_template("index.html")

@app.route("/analyze", methods=["POST"])
def analyze():
    job_desc = request.form.get("job_desc", "").strip()
    job_desc = re.sub(r'\s+', ' ', job_desc)
    uploaded = request.files.get("resume_file")

    if not uploaded or uploaded.filename == "":
        flash("Please upload a resume PDF or DOCX file.")
        return redirect(url_for("index"))

    tmp_fd, tmp_path = tempfile.mkstemp(suffix=os.path.splitext(uploaded.filename)[1] or ".pdf")
    uploaded.save(tmp_path)

    try:
        ext = uploaded.filename.lower()
        if ext.endswith(".pdf"):
            resume_text = extract_text_from_pdf(tmp_path)
        elif ext.endswith(".docx"):
            resume_text = extract_text_from_docx(tmp_path)
        else:
            try:
                with open(tmp_path, "r", encoding="utf-8") as f:
                    resume_text = f.read()
            except Exception:
                resume_text = extract_text_from_pdf(tmp_path)

        if not resume_text:
            resume_text = "(No text could be extracted from the uploaded file.)"

        before_score, matched_keywords, missing_keywords = compute_match_score(resume_text, job_desc)

        # Call Ollama
        ai_result = call_ollama_optimize(resume_text, job_desc)
        optimized_resume = ai_result.get("optimized_resume", resume_text)
        suggested_improvements = ai_result.get("suggested_improvements", [])
        roadmap = ai_result.get("roadmap", [])
        matched_from_model = ai_result.get("matched_keywords", matched_keywords)
        after_score, _, _ = compute_match_score(optimized_resume, job_desc)

        file_id = uuid.uuid4().hex
        out_path = GENERATED_DIR / f"{file_id}.docx"
        save_docx_from_text(optimized_resume, out_path)

    finally:
        try:
            os.close(tmp_fd)
            os.remove(tmp_path)
        except Exception:
            pass

    return render_template("results.html",
                           before_score=before_score,
                           after_score=after_score,
                           matched_keywords=matched_from_model or matched_keywords,
                           missing_keywords=missing_keywords,
                           suggested_improvements=suggested_improvements,
                           roadmap=roadmap,
                           optimized_resume=optimized_resume,
                           download_id=file_id
                           )

@app.route("/download/optimized/<file_id>", methods=["GET"])
def download_optimized(file_id):
    path = GENERATED_DIR / f"{file_id}.docx"
    if not path.exists():
        flash("File not found.")
        return redirect(url_for("index"))
    return send_file(path, as_attachment=True, download_name="Optimized_Resume.docx")

if __name__ == "__main__":
    app.run(debug=True, port=5000)
